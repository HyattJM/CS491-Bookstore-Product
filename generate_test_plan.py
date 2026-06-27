from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

document = Document()

# Add a title
document.add_heading('Test Plan: Rare Finds Bookstore Management System (BMS)', 0)

document.add_paragraph('Prepared by: [Your Name/Student ID]\nDate: June 27, 2026')

# 1. Test plan identifier
document.add_heading('1.0 Test Plan Identifier', level=1)
document.add_paragraph('TP-RFBMS-1.0')

# 2. Introduction
document.add_heading('2.0 Introduction', level=1)
document.add_paragraph('The Rare Finds Bookstore Management System (BMS) is a full-stack Web application designed to streamline inventory, sales, and secure user management. This test plan outlines the strategy, scope, and objectives for testing the Sprint 1 features, primarily focusing on Authentication, Role-Based Access Control (RBAC), CRUD operations for inventory, and Search & Filter UI functionality.')

# 3. Test objectives and tasks
document.add_heading('3.0 Test Objectives and Tasks', level=1)
document.add_paragraph('Objectives:\n- Validate that all features implemented in Sprint 1 function according to specifications.\n- Ensure Role-Based Access Control restricts functionality correctly for Admin, Manager, and Clerk roles.\n- Verify the accuracy of the Search & Filter feature.')
document.add_paragraph('Tasks:\n- Write and execute unit tests for backend services.\n- Perform integration testing on REST APIs.\n- Conduct system and UI testing on the React frontend.\n- Report and track any identified bugs or defects.')

# 4. Scope
document.add_heading('4.0 Scope', level=1)
document.add_paragraph('The scope of this testing encompasses the following features from Sprint 1:\n- User Authentication (Login) and Authorization (RBAC).\n- Inventory CRUD operations (Create, Read, Update, Delete books).\n- Search & Filter functionalities (by Title, Author, Genre).')
document.add_paragraph('Performance and stress testing are currently out of scope for Sprint 1 and will be addressed in future sprints.')

# 5. Test strategy
document.add_heading('5.0 Test Strategy', level=1)
document.add_paragraph('The testing strategy will incorporate a combination of manual and automated testing methods:')
document.add_paragraph('- Unit Testing: Developers will test individual components (e.g., Spring Data JPA queries) using JUnit.')
document.add_paragraph('- System and Integration Testing: API endpoints will be tested using Postman/Swagger to ensure correct responses and error handling.')
document.add_paragraph('- User Acceptance Testing (UAT): Manual testing of the React UI to verify user flows such as logging in, searching, and managing inventory.')

# 6. Features to be tested
document.add_heading('6.0 Features to be Tested', level=1)
document.add_paragraph('1. Authentication Module: Secure login using Base64 Basic Auth tokens.')
document.add_paragraph('2. Role-Based Access Control (RBAC): Admin and Manager can add books; only Admin can delete books.')
document.add_paragraph('3. Inventory Management: End-to-end functionality of adding, viewing, updating, and deleting book records.')
document.add_paragraph('4. Search and Filter UI: Keyword searches via the search bar and dropdown filter (Title, Author, Genre) returning accurate results.')

# 7. Features not to be tested
document.add_heading('7.0 Features Not to be Tested', level=1)
document.add_paragraph('Features planned for Sprint 2 are not included in this test plan:\n- Sales Processing\n- Reporting Dashboards\n- Supplier Management')

# 8. Environmental requirements
document.add_heading('8.0 Environmental Requirements', level=1)
document.add_paragraph('Hardware:\n- Standard workstation/laptop for developers and testers.')
document.add_paragraph('Software/Environment:\n- Backend: Java 21, Spring Boot 3.2.5, running on localhost:8081.\n- Database: MySQL 8.0+.\n- Frontend: React (SPA), Vite, Node.js.\n- Browsers: Latest versions of Google Chrome, Mozilla Firefox, and Microsoft Edge.')

# 9. Resources/roles and responsibilities
document.add_heading('9.0 Resources/Roles and Responsibilities', level=1)
document.add_paragraph('- Product Owner (Frank Gialluca): Approves test plan and accepts final release.')
document.add_paragraph('- Scrum Master (Scott Fleenor): Ensures testing adheres to the project schedule.')
document.add_paragraph('- Developers (James Hyatt, Yves Michel Fouty-Bikandou): Perform unit testing and fix reported defects.')
document.add_paragraph('- QA/Reporting (Luke Brown): Writes test cases, executes system testing, and reports defects.')

# 10. Test schedule
document.add_heading('10.0 Test Schedule', level=1)
document.add_paragraph('Test Planning: June 27, 2026\nTest Case Development: June 28, 2026\nTest Execution (Sprint 1): June 29 - July 1, 2026\nDefect Resolution and Retesting: July 2 - July 3, 2026\nFinal Sign-off: July 4, 2026')

# 11. Risks and mitigation
document.add_heading('11.0 Risks and Mitigation', level=1)
document.add_paragraph('Risk: Development delays push back the start of testing.\nMitigation: QA team will begin writing test cases in parallel with development using API specifications.')
document.add_paragraph('Risk: Testing environment is unstable or database is corrupted.\nMitigation: Automated database schema generation (schema.sql) and sample data seeding are implemented to quickly reset the environment.')

# 12. Approvals
document.add_heading('12.0 Approvals', level=1)
document.add_paragraph('Name: Frank Gialluca, Product Owner\nSignature: __________________________\nDate: ________________')

document.save('Rare_Finds_Test_Plan.docx')
